import pandas as pd
import google.generativeai as genai
import docx
import os
import threading # Para não travar a interface durante o processamento

# --- Bibliotecas para a Interface Gráfica ---
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext

# --- Funções Originais (com pequenas adaptações para a GUI) ---

def converter_excel_para_json_string(caminho_excel: str, nome_planilha: str | int = 0, status_log=None) -> str | None:
    if status_log: status_log(f"📄 Lendo o arquivo Excel: '{os.path.basename(caminho_excel)}'...")
    try:
        df = pd.read_excel(caminho_excel, sheet_name=nome_planilha)
        json_string = df.to_json(orient='records', indent=2, force_ascii=False)
        if status_log: status_log("✅ Dados do Excel convertidos para JSON em memória.")
        return json_string
    except FileNotFoundError:
        messagebox.showerror("Erro de Arquivo", f"O arquivo '{caminho_excel}' não foi encontrado.")
        return None
    except Exception as e:
        messagebox.showerror("Erro de Leitura", f"❌ ERRO ao ler o Excel: {e}")
        return None

def analisar_dados_com_gemini(json_data: str, prompt_usuario: str, status_log=None) -> str | None:
    if status_log: status_log("🤖 Conectando à API do Gemini para análise...")
    try:
        with open('config.txt', 'r', encoding='utf-8') as f:
            api_key = f.read().strip()
        if not api_key:
            messagebox.showerror("Erro de API", "O arquivo 'config.txt' está vazio. Cole sua chave da API nele.")
            return None
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt_completo = f"{prompt_usuario}\n\nAqui estão os dados em formato JSON para sua análise:\n\n{json_data}"
        response = model.generate_content(prompt_completo)
        if status_log: status_log("✅ Análise recebida do Gemini.")
        return response.text
    except FileNotFoundError:
        messagebox.showerror("Erro de Configuração", "Arquivo 'config.txt' não encontrado! Crie este arquivo na mesma pasta do .exe e cole sua chave da API do Google nele.")
        return None
    except Exception as e:
        messagebox.showerror("Erro de API", f"❌ ERRO durante a chamada da API do Gemini: {e}")
        return None

def salvar_texto_em_docx(texto: str, caminho_docx: str, status_log=None):
    if status_log: status_log(f"💾 Salvando a análise no documento Word: '{os.path.basename(caminho_docx)}'...")
    try:
        documento = docx.Document()
        documento.add_heading('Análise de Dados Gerada por IA', level=1)
        documento.add_paragraph(texto)
        documento.save(caminho_docx)
        if status_log: status_log(f"🎉 Documento '{os.path.basename(caminho_docx)}' salvo com sucesso!")
        messagebox.showinfo("Sucesso!", f"Relatório salvo com sucesso em:\n{caminho_docx}")
    except Exception as e:
        messagebox.showerror("Erro ao Salvar", f"❌ ERRO ao salvar o arquivo .docx: {e}")



def iniciar_processamento(caminho_excel, caminho_saida, status_log):
    """Função que executa a lógica principal em uma thread separada."""
    try:
        # --- MODIFICAÇÃO AQUI ---
        # Tenta ler o prompt do arquivo externo.
        status_log("📝 Lendo o prompt do arquivo 'prompt.txt'...")
        try:
            with open('prompt.txt', 'r', encoding='utf-8') as f:
                PROMPT_ANALISE = f.read()
            status_log("✅ Prompt carregado com sucesso.")
        except FileNotFoundError:
            messagebox.showerror("Erro de Configuração", "Arquivo 'prompt.txt' não encontrado! Certifique-se de que ele está na mesma pasta do executável.")
            status_log("🔴 ERRO: 'prompt.txt' não encontrado.")
            return # Interrompe a execução se o prompt não for encontrado
        # --- FIM DA MODIFICAÇÃO ---
        
        json_dos_dados = converter_excel_para_json_string(caminho_excel, status_log=status_log)

        if json_dos_dados:
            analise_gemini = analisar_dados_com_gemini(json_dos_dados, PROMPT_ANALISE, status_log=status_log)
            
            if analise_gemini:
                salvar_texto_em_docx(analise_gemini, caminho_saida, status_log=status_log)
            else:
                if status_log: status_log("\n🔴 Processo interrompido devido a erro na análise.")
        else:
            if status_log: status_log("\n🔴 Processo interrompido devido a erro na leitura do Excel.")
            
    except Exception as e:
        messagebox.showerror("Erro Inesperado", f"Ocorreu um erro inesperado: {e}")
    finally:
        # Reabilita o botão após o término do processo
        app.start_button.config(state=tk.NORMAL)



# --- Classe da Interface Gráfica (GUI) ---
class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Analisador de Dados com IA")
        self.root.geometry("600x450")

        self.caminho_excel = ""
        self.caminho_saida = ""

        # --- Widgets ---
        frame = tk.Frame(root, padx=10, pady=10)
        frame.pack(fill=tk.BOTH, expand=True)

        # Seleção de arquivo de entrada
        btn_select_excel = tk.Button(frame, text="1. Selecionar Arquivo Excel (.xlsx)", command=self.selecionar_excel)
        btn_select_excel.pack(fill=tk.X, pady=5)
        self.lbl_excel = tk.Label(frame, text="Nenhum arquivo selecionado", fg="gray")
        self.lbl_excel.pack()
        
        # Seleção de local de saída
        btn_select_output = tk.Button(frame, text="2. Definir Local de Saída (.docx)", command=self.definir_saida)
        btn_select_output.pack(fill=tk.X, pady=5)
        self.lbl_output = tk.Label(frame, text="Nenhum local de saída definido", fg="gray")
        self.lbl_output.pack()

        # Botão de Iniciar
        self.start_button = tk.Button(frame, text="🚀 INICIAR ANÁLISE", bg="#4CAF50", fg="white", font=('Helvetica', 10, 'bold'), command=self.iniciar)
        self.start_button.pack(fill=tk.X, pady=(20, 10), ipady=5)

        # Área de Log/Status
        self.log_area = scrolledtext.ScrolledText(frame, height=10, state='disabled', bg="#f0f0f0")
        self.log_area.pack(fill=tk.BOTH, expand=True, pady=(10, 0))

    def log(self, message):
        """Adiciona mensagens à área de log na interface."""
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, message + "\n")
        self.log_area.see(tk.END) # Auto-scroll
        self.log_area.config(state='disabled')
        self.root.update_idletasks() # Força a atualização da UI

    def selecionar_excel(self):
        self.caminho_excel = filedialog.askopenfilename(
            title="Selecione o arquivo Excel",
            filetypes=(("Arquivos Excel", "*.xlsx"), ("Todos os arquivos", "*.*"))
        )
        if self.caminho_excel:
            self.lbl_excel.config(text=os.path.basename(self.caminho_excel), fg="black")

    def definir_saida(self):
        self.caminho_saida = filedialog.asksaveasfilename(
            title="Salvar relatório como...",
            defaultextension=".docx",
            filetypes=(("Documento Word", "*.docx"), ("Todos os arquivos", "*.*"))
        )
        if self.caminho_saida:
            self.lbl_output.config(text=os.path.basename(self.caminho_saida), fg="black")

    def iniciar(self):
        if not self.caminho_excel or not self.caminho_saida:
            messagebox.showwarning("Atenção", "Por favor, selecione o arquivo de entrada e o local de saída antes de iniciar.")
            return

        self.log_area.config(state='normal')
        self.log_area.delete(1.0, tk.END) # Limpa o log
        self.log_area.config(state='disabled')
        
        self.start_button.config(state=tk.DISABLED) # Desabilita o botão para evitar cliques duplos
        
        # Inicia o processamento em uma nova thread para não congelar a GUI
        thread = threading.Thread(target=iniciar_processamento, args=(self.caminho_excel, self.caminho_saida, self.log))
        thread.start()

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()